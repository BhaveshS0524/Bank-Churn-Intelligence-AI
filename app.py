import streamlit as st
import pandas as pd
import plotly.express as px
import google.generativeai as genai
import plotly.io as pio

pio.templates.default = "plotly_white"

# ---------------- CONFIG ----------------
st.set_page_config(page_title="BFSI Churn Intelligence", layout="wide")

# ---------------- LOAD DATA ----------------
@st.cache_data
def load_data():
    df = pd.read_csv("Bank_Churn.csv")
    return df

df = load_data()

# ---------------- FEATURE ENGINEERING ----------------
df["CustomerValue"] = df["Balance"] * df["NumOfProducts"]
df["EngagementScore"] = df["IsActiveMember"] * df["NumOfProducts"]
df["RevenueRisk"] = df["Balance"] * df["Exited"]

# ---------------- HEADER ----------------
st.title("🏦 BFSI Customer Churn Intelligence Platform")
st.markdown("### AI-Powered Retention & Revenue Risk System")

# ---------------- KPI METRICS ----------------
total_customers = len(df)
churned = df["Exited"].sum()
churn_rate = (churned / total_customers) * 100
total_revenue_risk = df["RevenueRisk"].sum()

col1, col2, col3, col4 = st.columns(4)
col1.metric("Total Customers", f"{total_customers:,}")
col2.metric("Churn Rate", f"{churn_rate:.2f}%")
col3.metric("Churned Customers", f"{churned:,}")
col4.metric("Revenue at Risk ($)", f"{total_revenue_risk:,.0f}")

st.divider()

# ---------------- SEGMENTATION ----------------
st.subheader("🎯 Customer Risk Segmentation")

def segment_customer(row):
    if row["Exited"] == 1 and row["Balance"] > 100000:
        return "High Value - High Risk"
    elif row["Exited"] == 1:
        return "Low Value - High Risk"
    elif row["IsActiveMember"] == 1:
        return "Loyal"
    else:
        return "At Risk"

df["Segment"] = df.apply(segment_customer, axis=1)

seg_counts = df["Segment"].value_counts().reset_index()
fig_seg = px.pie(seg_counts, names="Segment", values="count",
                 title="Customer Segmentation")
st.plotly_chart(fig_seg, use_container_width=True)

# ---------------- VISUALS ----------------
geo = df.groupby("Geography")["Exited"].mean().reset_index()
geo["Exited"] *= 100
geo = geo.sort_values(by="Exited")

fig_geo = px.bar(
    geo,
    x="Exited",
    y="Geography",
    orientation="h",
    text=geo["Exited"].round(1),
    color="Exited",
    color_continuous_scale=[
        "#fde0dd", "#fa9fb5", "#c51b8a", "#7a0177"
    ]
)

fig_geo.update_traces(
    textposition="outside",
    textfont=dict(size=14, color="black"),
    cliponaxis=False   # 🔥 KEY FIX
)

fig_geo.update_layout(
    title="Churn Rate by Geography (%)",
    xaxis_title="Churn Rate (%)",
    yaxis_title="",
    font=dict(size=14),
    margin=dict(l=40, r=80, t=60, b=40)
)

st.plotly_chart(fig_geo, use_container_width=True)

# 🔥 2. Balance Distribution (Business-Friendly)
with col2:
    st.subheader("Balance Distribution vs Churn")

    fig_balance = px.histogram(
    df,
    x="Balance",
    color="Exited",
    nbins=40,
    opacity=0.75,
    barmode="overlay",
    color_discrete_map={
        0: "#27ae60",   # strong green
        1: "#c0392b"    # strong red
    }
)

fig_balance.update_layout(
    title="Balance Distribution: Churn vs Retained",
    xaxis_title="Account Balance",
    yaxis_title="Customers",
    font=dict(size=14),
)

st.plotly_chart(fig_balance, use_container_width=True)

# 🔥 3. Customer Segmentation (Donut Chart)
st.subheader("Customer Segmentation")

fig_seg = px.pie(
    seg_counts,
    names="Segment",
    values="count",
    hole=0.55,
    color_discrete_sequence=[
        "#e74c3c", "#f39c12", "#2ecc71", "#3498db"
    ]
)

fig_seg.update_traces(
    textinfo="percent+label",
    textfont=dict(size=13)
)

st.plotly_chart(fig_seg, use_container_width=True)

# ---------------- TOP RISK CUSTOMERS ----------------
st.subheader("⚠️ High Revenue Risk Customers")

top_risk = df[df["Exited"] == 1].sort_values(by="Balance", ascending=False).head(10)
st.dataframe(top_risk[["CustomerId", "Balance", "Geography", "NumOfProducts"]])


st.divider()
st.header("🧠 AI Decision Intelligence Layer")

# 1. Risk Alerts
st.subheader("⚠️ Risk Alerts")
# st.subheader("⚠️ Risk Alerts")

high_risk_customers = df[
    (df["Balance"] > 100000) &
    (df["IsActiveMember"] == 0) &
    (df["Exited"] == 1)
]

if len(high_risk_customers) > 0:
    st.error(f"🚨 {len(high_risk_customers)} high-value customers have churned!")
else:
    st.success("✅ No critical churn risk detected")

# 2. Top Customers to Save
st.subheader("🔮 Top Customers to Save")
# st.subheader("🔮 Top Customers to Save")

df["ChurnProbability"] = (
    (1 - df["IsActiveMember"]) * 0.4 +
    (df["Balance"] / df["Balance"].max()) * 0.4 +
    (df["NumOfProducts"] <= 2) * 0.2
)

top_save = df.sort_values(
    by="ChurnProbability",
    ascending=False
).head(10)

st.dataframe(
    top_save[[
        "CustomerId",
        "Balance",
        "Geography",
        "NumOfProducts",
        "ChurnProbability"
    ]]
)

# 3. Ask AI

st.divider()
st.header("❓ Frequently Asked Questions (Ask AI)")

# ---------- BUSINESS INSIGHTS ----------
with st.expander("📊 Business Insights"):
    st.markdown("""
- Which customers are most likely to churn and why?
- What are the main drivers of churn in this dataset?
- Which geography has the highest churn risk and what could be the reason?
- What trends do you observe in customer churn behavior?
""")

# ---------- REVENUE & RISK ----------
with st.expander("💰 Revenue & Risk Analysis"):
    st.markdown("""
- What is the estimated revenue at risk due to churn?
- Which customer segment contributes the most to revenue loss?
- How can the bank reduce financial impact from churn?
- Which high-value customers should be prioritized for retention?
""")

# ---------- RETENTION STRATEGY ----------
with st.expander("🎯 Retention Strategy"):
    st.markdown("""
- What retention strategies should be applied to high-risk customers?
- Suggest personalized offers to retain premium customers
- How can we improve engagement for inactive users?
- What actions should be taken immediately to reduce churn?
""")

# ---------- DATA ANALYSIS ----------
with st.expander("📈 Data Analysis & Patterns"):
    st.markdown("""
- What is the relationship between balance and churn?
- How does customer activity impact churn?
- Which age group has the highest churn rate?
- Does the number of products affect retention?
""")

# ---------- SCENARIO ANALYSIS ----------
with st.expander("🔮 Scenario Analysis"):
    st.markdown("""
- What will happen if customer engagement increases by 20%?
- How would reducing churn by 5% impact revenue?
- What if high-balance customers become inactive?
""")

# ---------- USER GUIDANCE ----------
st.info("💡 Tip: Ask business-focused questions like 'How can we reduce churn?' for better AI insights.")

if st.button("Which customers are most likely to churn and why?"):
    user_query = "Which customers are most likely to churn and why?"

# ---------------- ASK AI ----------------
st.subheader("🧠 Ask AI About Customer Data")

user_query = st.text_input("Ask a business question:")

if user_query:
    summary = f"""
    Dataset Summary:
    - Total Customers: {len(df)}
    - Churn Rate: {round(df['Exited'].mean()*100,2)}%
    - Avg Balance: {round(df['Balance'].mean(),2)}
    """

    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
    model = genai.GenerativeModel("gemini-2.5-flash")

    prompt = f"""
    You are a senior banking data analyst.

    {summary}

    User Question:
    {user_query}

    Provide:
    - Clear business insight
    - Data-backed reasoning
    - Actionable recommendation
    """

    with st.spinner("Analyzing..."):
        try:
            response = model.generate_content(prompt)

            # ✅ Show AI response
            st.markdown("### 🤖 AI Insight")
            st.write(response.text)

            # ---------------- DOWNLOAD SECTION ----------------
            from docx import Document
            from io import BytesIO

            def create_docx(report_text):
                doc = Document()
                doc.add_heading("Customer Churn Intelligence Report", 0)
                doc.add_paragraph(report_text)

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer

            docx_file = create_docx(response.text)

            st.markdown("### 📥 Export Report")

            col1, col2 = st.columns(2)

            with col1:
                st.download_button(
                    "📄 Download TXT",
                    response.text,
                    file_name="churn_report.txt",
                    mime="text/plain"
                )

            with col2:
                st.download_button(
                    "📄 Download DOCX",
                    docx_file,
                    file_name="churn_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except:
            st.error("AI service unavailable")

# ---------------- SIDEBAR INPUT ----------------
st.sidebar.header("🔍 Customer 360 Analysis")

cs = st.sidebar.slider("Credit Score", 300, 850, 650)
geo = st.sidebar.selectbox("Geography", ["France", "Germany", "Spain"])
age = st.sidebar.number_input("Age", 18, 100, 35)
tenure = st.sidebar.slider("Tenure", 0, 10, 5)
balance = st.sidebar.number_input("Balance", value=50000)
products = st.sidebar.selectbox("Products", [1, 2, 3, 4])
active = st.sidebar.radio("Active Member", ["Yes", "No"])

analyze = st.sidebar.button("🚀 Analyze Customer")

# ---------------- BUSINESS LOGIC ----------------
def calculate_risk(balance, age, active, products):
    active_val = 1 if active == "Yes" else 0
    
    if balance > 100000 and active_val == 0:
        return "High"
    elif age > 45 and products <= 2:
        return "Medium"
    else:
        return "Low"

# ---------------- AI ENGINE ----------------
if analyze:
    risk = calculate_risk(balance, age, active, products)
    revenue_risk = balance if risk == "High" else balance * 0.5
    
    st.divider()
    st.subheader(f"🧠 Customer Intelligence Report ({risk} Risk)")

    col1, col2 = st.columns(2)
    col1.metric("Estimated Revenue Risk", f"${revenue_risk:,.0f}")
    col2.metric("Risk Level", risk)

    # Gemini Config
    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
    model = genai.GenerativeModel("gemini-2.5-flash")

    prompt = f"""
    You are a Senior Banking Retention Strategist.

    Customer Profile:
    - Age: {age}
    - Geography: {geo}
    - Balance: {balance}
    - Products: {products}
    - Active: {active}

    Risk Level: {risk}

    Provide:
    1. Key churn drivers
    2. Revenue impact
    3. 3-step retention strategy
    4. Suggested personalized offer
    """

    with st.spinner("Generating AI strategy..."):
        try:
            response = model.generate_content(prompt)
            st.write(response.text)
        except:
            st.error("AI service unavailable")

st.write(response.text)

# ---------------- ASK AI ----------------
st.subheader("🧠 Ask AI About Customer Data")

user_query = st.text_input("Ask a business question:")

if user_query:
    summary = f"""
    Dataset Summary:
    - Total Customers: {len(df)}
    - Churn Rate: {round(df['Exited'].mean()*100,2)}%
    - Avg Balance: {round(df['Balance'].mean(),2)}
    """

    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
    model = genai.GenerativeModel("gemini-2.5-flash")

    prompt = f"""
    You are a senior banking data analyst.

    {summary}

    User Question:
    {user_query}

    Provide:
    - Clear business insight
    - Data-backed reasoning
    - Actionable recommendation
    """

    with st.spinner("Analyzing..."):
        try:
            response = model.generate_content(prompt)

            # ✅ Show AI response
            st.markdown("### 🤖 AI Insight")
            st.write(response.text)

            # ---------------- DOWNLOAD SECTION ----------------
            from docx import Document
            from io import BytesIO

            from docx import Document
from io import BytesIO

def create_enterprise_docx(report_text, user_query):
    doc = Document()

    # Title
    doc.add_heading("BFSI Customer Churn Intelligence Report", 0)

    # Subtitle
    doc.add_paragraph("AI-Powered Strategic Analysis for Banking Decision Makers\n")

    # Query
    doc.add_heading("User Query", level=1)
    doc.add_paragraph(user_query)

    # AI Insights
    doc.add_heading("AI-Generated Insights", level=1)
    doc.add_paragraph(report_text)

    # Recommendations section
    doc.add_heading("Recommended Actions", level=1)
    doc.add_paragraph(
        "- Prioritize high-value customers\n"
        "- Improve engagement strategies\n"
        "- Monitor churn-prone segments"
    )

    # Footer
    doc.add_paragraph("\nGenerated using AI Churn Intelligence Platform")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from io import BytesIO

def create_pdf(report_text):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()

    content = []

    content.append(Paragraph("BFSI Customer Churn Intelligence Report", styles["Title"]))
    content.append(Paragraph("<br/>", styles["Normal"]))
    content.append(Paragraph(report_text, styles["Normal"]))

    doc.build(content)
    buffer.seek(0)

    return buffer

# ---------------- EXPORT SECTION ----------------

st.markdown("### 📥 Export Enterprise Report")

docx_file = create_enterprise_docx(response.text, user_query)
pdf_file = create_pdf(response.text)

col1, col2, col3 = st.columns(3)

with col1:
    st.download_button(
        "📄 Download DOCX",
        docx_file,
        file_name="BFSI_Churn_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

with col2:
    st.download_button(
        "📕 Download PDF",
        pdf_file,
        file_name="BFSI_Churn_Report.pdf",
        mime="application/pdf"
    )

with col3:
    st.download_button(
        "📄 Download TXT",
        response.text,
        file_name="BFSI_Report.txt"
    )

st.markdown("### ✉️ Copy for Email")

email_text = f"""
Subject: Customer Churn Risk Analysis

Dear Team,

Please find below the key insights:

{response.text}

Regards,  
Churn Intelligence System
"""

st.text_area("Email Draft", email_text, height=200)

# ---------------- WHAT-IF SIMULATION ----------------
st.divider()
st.subheader("🔮 Scenario Simulation")

new_balance = st.slider("Adjust Balance", 0, 200000, 50000)

sim_risk = calculate_risk(new_balance, age, active, products)

st.metric("Simulated Risk Level", sim_risk)

# ---------------- FOOTER ----------------
st.sidebar.divider()
st.sidebar.markdown("### 💼 About")
st.sidebar.info("AI-powered BFSI churn intelligence system for decision-makers.")

# --- STEP 6: PORTFOLIO & SOCIAL SHARING ---
st.sidebar.divider()
st.sidebar.subheader("📢 Share this Insight")
st.sidebar.info("Found an interesting risk profile? Share this tool with your network.")

# Replace with your actual LinkedIn profile link
linkedin_url = "https://www.linkedin.com/in/bhaveshsuryavanshi/" 
st.sidebar.link_button("🤝 Connect with the Developer", linkedin_url)

# Add a "Copy Link" helper
st.sidebar.code(f"App Link: {st.secrets.get('APP_URL', 'bank-churn-intelligence-bhavesh.streamlit.app')}")